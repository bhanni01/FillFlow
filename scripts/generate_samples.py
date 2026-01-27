"""
Generate sample-template.docx and sample-data.xlsx for Word Data Merger testing.

sample-template.docx uses python-docx to insert real Word content controls
(Plain Text Content Controls) with tags.

sample-data.xlsx uses openpyxl with realistic invoice data.
"""

import os
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from datetime import date

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "samples")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ─── Helpers: build a Word content control (OOXML) ───────────────────────────

def make_sdt(tag: str, placeholder: str = None) -> OxmlElement:
    """
    Build an inline <w:sdt> (Plain Text Content Control) with a given tag.
    Returns the sdt element to be inserted into a paragraph's XML.
    """
    sdt = OxmlElement("w:sdt")

    # sdtPr — properties
    sdtPr = OxmlElement("w:sdtPr")

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), tag)
    sdtPr.append(alias)

    tagEl = OxmlElement("w:tag")
    tagEl.set(qn("w:val"), tag)
    sdtPr.append(tagEl)

    # Lock content (optional — makes it look cleaner)
    lock = OxmlElement("w:lock")
    lock.set(qn("w:val"), "sdtContentLocked")
    sdtPr.append(lock)

    # Plain text type
    text_type = OxmlElement("w:text")
    sdtPr.append(text_type)

    # Show placeholder text
    if placeholder:
        ph = OxmlElement("w:showingPlcHdr")
        sdtPr.append(ph)
        sdtContent_ph = OxmlElement("w:sdtPr")  # won't use this path

    sdt.append(sdtPr)

    # sdtContent — holds the visible content
    sdtContent = OxmlElement("w:sdtContent")
    p_inner = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = placeholder or tag
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p_inner.append(r)
    sdtContent.append(p_inner)
    sdt.append(sdtContent)

    return sdt


def make_inline_sdt(tag: str, placeholder: str = None) -> OxmlElement:
    """
    Build an inline (run-level) <w:sdt> for use inside a paragraph
    (not as a block-level SDT).
    """
    sdt = OxmlElement("w:sdt")

    sdtPr = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), tag)
    sdtPr.append(alias)

    tagEl = OxmlElement("w:tag")
    tagEl.set(qn("w:val"), tag)
    sdtPr.append(tagEl)

    text_type = OxmlElement("w:text")
    sdtPr.append(text_type)

    sdt.append(sdtPr)

    sdtContent = OxmlElement("w:sdtContent")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = placeholder or tag
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    sdtContent.append(r)
    sdt.append(sdtContent)

    return sdt


def append_label_and_sdt(para, label: str, tag: str, placeholder: str):
    """Add a label run and an inline SDT to an existing paragraph."""
    # Label run
    label_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    b = OxmlElement("w:b")
    rPr.append(b)
    label_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = label
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    label_run.append(t)
    para._p.append(label_run)

    # Inline SDT
    sdt = make_inline_sdt(tag, placeholder)
    para._p.append(sdt)


# ─── Build sample-template.docx ──────────────────────────────────────────────

def build_template():
    doc = Document()

    # ── Page title ──
    title = doc.add_paragraph()
    title.alignment = 1  # center
    run = title.add_run("INVOICE")
    run.bold = True
    run.font.size = Pt(22)

    doc.add_paragraph()  # spacer

    # ── Flat-field block ──
    # Each field on its own line using inline SDTs
    fields = [
        ("Client:         ", "client_name",     "Acme Corp"),
        ("Invoice #:      ", "invoice_number",  "INV-0001"),
        ("Invoice Date:   ", "invoice_date",    "2026-06-18"),
        ("Total Due:      ", "total_due",       "$0.00"),
    ]

    for label, tag, placeholder in fields:
        para = doc.add_paragraph()
        append_label_and_sdt(para, label, tag, placeholder)

    doc.add_paragraph()  # spacer

    # ── Line-item table ──
    headers = ["Description", "Qty", "Unit Price", "Line Total"]
    item_tags = ["item_description", "item_qty", "item_unit_price", "item_total"]

    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"

    # Header row
    hdr_row = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(hdr)
        run.bold = True
        # Grey fill
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)

    # Template row (last body row) — each cell gets an inline SDT
    tmpl_row = tbl.rows[1]
    for i, tag in enumerate(item_tags):
        cell = tmpl_row.cells[i]
        # Clear default empty paragraph content
        para = cell.paragraphs[0]
        # Remove any existing runs
        for child in list(para._p):
            para._p.remove(child)
        # Add inline SDT
        sdt = make_inline_sdt(tag, tag)
        para._p.append(sdt)

    doc.add_paragraph()  # spacer after table

    footer_p = doc.add_paragraph("Thank you for your business.")
    footer_p.alignment = 1

    out_path = os.path.join(OUTPUT_DIR, "sample-template.docx")
    doc.save(out_path)
    print(f"Created: {out_path}")


# ─── Build sample-data.xlsx ──────────────────────────────────────────────────

def build_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Invoice Data"

    # ── Column headers ──
    headers = [
        "client_name", "invoice_number", "invoice_date", "total_due",
        "item_description", "item_qty", "item_unit_price", "item_total"
    ]
    ws.append(headers)

    # Style header row
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    # ── Row 2: flat invoice data (line-item columns left blank) ──
    ws.append([
        "Globex Corporation",     # client_name
        "INV-2026-0042",          # invoice_number
        "2026-06-18",             # invoice_date
        "$4,725.00",              # total_due
        "", "", "", ""            # item columns blank for flat row
    ])

    # ── Rows 3–7: line items (flat columns blank) ──
    line_items = [
        ("Strategy Consulting — 10 hrs", 10, 250.00, 2500.00),
        ("Market Research Report",        1, 850.00,  850.00),
        ("Brand Identity Package",        1, 600.00,  600.00),
        ("Social Media Setup",            3,  95.00,  285.00),
        ("Project Management Fee",        1, 490.00,  490.00),
    ]

    for desc, qty, unit, total in line_items:
        ws.append([
            "", "", "", "",         # flat columns blank for item rows
            desc,
            qty,
            f"${unit:,.2f}",
            f"${total:,.2f}"
        ])

    # ── Column widths ──
    col_widths = [22, 16, 14, 12, 34, 8, 14, 14]
    for i, w in enumerate(col_widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    # Alternate row fill for readability
    light_fill = PatternFill(start_color="EBF3FB", end_color="EBF3FB", fill_type="solid")
    for row_idx in range(3, ws.max_row + 1, 2):
        for cell in ws[row_idx]:
            cell.fill = light_fill

    out_path = os.path.join(OUTPUT_DIR, "sample-data.xlsx")
    wb.save(out_path)
    print(f"Created: {out_path}")


if __name__ == "__main__":
    build_template()
    build_xlsx()
    print("Done.")
